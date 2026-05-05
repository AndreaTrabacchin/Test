#!/usr/bin/env python3
"""
Follow-ups Parser Script
Converts Word files, CSV, or plain text into JSON format for the follow-ups management system.

Usage:
    python parse_followups.py --input followups.docx --output followups.json
    python parse_followups.py --input followups.csv --output followups.json
    python parse_followups.py --input followups.txt --output followups.json
"""

import json
import csv
import re
import argparse
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any

try:
    from docx import Document
except ImportError:
    Document = None


class FollowupsParser:
    """Parse follow-ups from various file formats"""
    
    def __init__(self):
        self.follow_ups = []
        self.id_counter = 1
    
    def generate_id(self) -> str:
        """Generate unique follow-up ID"""
        id_str = f"FU-{str(self.id_counter).zfill(3)}"
        self.id_counter += 1
        return id_str
    
    def parse_word_file(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse follow-ups from Word document"""
        if Document is None:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        doc = Document(file_path)
        follow_ups = []
        
        # Try to parse from table
        if doc.tables:
            follow_ups = self._parse_from_table(doc.tables[0])
        
        # Fall back to parsing from paragraphs
        if not follow_ups:
            follow_ups = self._parse_from_paragraphs(doc.paragraphs)
        
        return follow_ups
    
    def _parse_from_table(self, table) -> List[Dict[str, Any]]:
        """Extract follow-ups from Word table"""
        follow_ups = []
        
        # Skip header row if present
        start_row = 1 if len(table.rows) > 1 else 0
        
        for row in table.rows[start_row:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 3:
                follow_up = {
                    "id": self.generate_id(),
                    "action": cells[0],
                    "owner_name": cells[1],
                    "owner_email": cells[2] if len(cells) > 2 else "",
                    "expected_date": cells[3] if len(cells) > 3 else "",
                    "status": "pending",
                    "created_date": datetime.now().strftime("%Y-%m-%d")
                }
                follow_ups.append(follow_up)
        
        return follow_ups
    
    def _parse_from_paragraphs(self, paragraphs) -> List[Dict[str, Any]]:
        """Extract follow-ups from Word paragraphs using pattern matching"""
        follow_ups = []
        
        for para in paragraphs:
            text = para.text.strip()
            if not text or text.startswith("Follow-ups") or text.startswith("#"):
                continue
            
            # Look for patterns like "Action: ... Owner: ... Date: ..."
            follow_up = self._extract_from_text(text)
            if follow_up:
                follow_ups.append(follow_up)
        
        return follow_ups
    
    def _extract_from_text(self, text: str) -> Dict[str, Any]:
        """Extract follow-up from free-form text"""
        patterns = {
            "action": r"(?:Action|Task|Follow-?up):\s*(.+?)(?:Owner|Contact|To:|\nDate|\n|$)",
            "owner_name": r"(?:Owner|Contact|To):\s*(.+?)(?:\(|<|Owner Email|Email:|\n|$)",
            "owner_email": r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
            "expected_date": r"(?:Date|Expected|Deadline):\s*(\d{4}-\d{2}-\d{2}|\d{2}/\d{2}/\d{4})"
        }
        
        follow_up = {
            "id": self.generate_id(),
            "status": "pending",
            "created_date": datetime.now().strftime("%Y-%m-%d")
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                value = match.group(1).strip()
                if key == "expected_date":
                    value = self._normalize_date(value)
                follow_up[key] = value
            else:
                follow_up[key] = ""
        
        # If we found at least action and owner, it's valid
        if follow_up.get("action") and follow_up.get("owner_name"):
            return follow_up
        
        return None
    
    def parse_csv_file(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse follow-ups from CSV file"""
        follow_ups = []
        
        with open(file_path, 'r', encoding='utf-8') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                follow_up = {
                    "id": self.generate_id(),
                    "action": row.get('action') or row.get('Action') or row.get('Task') or "",
                    "owner_name": row.get('owner_name') or row.get('Owner') or row.get('Contact') or "",
                    "owner_email": row.get('owner_email') or row.get('Email') or row.get('email') or "",
                    "expected_date": self._normalize_date(row.get('expected_date') or row.get('Date') or row.get('date') or ""),
                    "status": "pending",
                    "created_date": datetime.now().strftime("%Y-%m-%d")
                }
                if follow_up["action"] and follow_up["owner_name"]:
                    follow_ups.append(follow_up)
        
        return follow_ups
    
    def parse_text_file(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse follow-ups from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        follow_ups = []
        
        # Split by common delimiters (---, ===, or multiple newlines)
        sections = re.split(r'---+|===+|\n\n+', content)
        
        for section in sections:
            if section.strip():
                follow_up = self._extract_from_text(section)
                if follow_up:
                    follow_ups.append(follow_up)
        
        return follow_ups
    
    def _normalize_date(self, date_str: str) -> str:
        """Normalize date to YYYY-MM-DD format"""
        if not date_str:
            return ""
        
        date_str = date_str.strip()
        
        # Try common formats
        formats = ["%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y", "%m-%d-%Y"]
        
        for fmt in formats:
            try:
                parsed_date = datetime.strptime(date_str, fmt)
                return parsed_date.strftime("%Y-%m-%d")
            except ValueError:
                continue
        
        # If no format matched, return as-is
        return date_str
    
    def to_json(self, follow_ups: List[Dict[str, Any]]) -> str:
        """Convert follow-ups to JSON format"""
        data = {"follow_ups": follow_ups}
        return json.dumps(data, indent=2)
    
    def save_json(self, follow_ups: List[Dict[str, Any]], output_path: str):
        """Save follow-ups to JSON file"""
        json_str = self.to_json(follow_ups)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(json_str)
        
        print(f"✅ Successfully saved {len(follow_ups)} follow-ups to {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Convert follow-ups from Word, CSV, or TXT to JSON format"
    )
    parser.add_argument(
        "--input", "-i",
        required=True,
        help="Input file path (docx, csv, or txt)"
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Output JSON file path (default: input_name_converted.json)"
    )
    parser.add_argument(
        "--print", "-p",
        action="store_true",
        help="Print output to console instead of saving"
    )
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    
    if not input_path.exists():
        print(f"❌ Error: Input file not found: {args.input}")
        return
    
    file_ext = input_path.suffix.lower()
    
    parser_instance = FollowupsParser()
    
    try:
        if file_ext == ".docx":
            print(f"📄 Parsing Word document: {args.input}")
            follow_ups = parser_instance.parse_word_file(args.input)
        elif file_ext == ".csv":
            print(f"📊 Parsing CSV file: {args.input}")
            follow_ups = parser_instance.parse_csv_file(args.input)
        elif file_ext == ".txt":
            print(f"📝 Parsing text file: {args.input}")
            follow_ups = parser_instance.parse_text_file(args.input)
        else:
            print(f"❌ Error: Unsupported file format: {file_ext}")
            print("   Supported formats: .docx, .csv, .txt")
            return
        
        if not follow_ups:
            print("⚠️  Warning: No follow-ups found in the input file")
            return
        
        print(f"✨ Found {len(follow_ups)} follow-ups\n")
        
        if args.print:
            print(parser_instance.to_json(follow_ups))
        else:
            output_path = args.output or f"{input_path.stem}_converted.json"
            parser_instance.save_json(follow_ups, output_path)
    
    except Exception as e:
        print(f"❌ Error: {e}")


if __name__ == "__main__":
    main()
